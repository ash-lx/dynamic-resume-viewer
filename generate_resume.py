from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import subprocess
import os


def create_resume_template(resume_data):
    """
    Creates a resume document based on a flexible template and provided data.
    """
    document = Document()

    # Define Custom Styles
    styles = document.styles

    # Heading 1 Style
    h1_style = styles.add_style('H1Custom', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.base_style = styles['Heading 1']
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(20)
    h1_style.font.bold = True
    h1_style.paragraph_format.space_after = Pt(8)
    h1_style.paragraph_format.keep_with_next = True

    # Heading 2 Style
    h2_style = styles.add_style('H2Custom', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.base_style = styles['Heading 2']
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(16)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x41, 0x69, 0xE1)
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(4)
    h2_style.paragraph_format.keep_with_next = True

    # Heading 3 Style
    h3_style = styles.add_style('H3Custom', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.base_style = styles['Heading 3']
    h3_style.font.name = 'Arial'
    h3_style.font.size = Pt(14)
    h3_style.font.bold = True
    h3_style.paragraph_format.space_before = Pt(10)
    h3_style.paragraph_format.space_after = Pt(3)

    # Normal Style
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)

    # List Bullet Style
    list_bullet_style = styles['List Bullet']
    list_bullet_style.font.name = 'Calibri'
    list_bullet_style.font.size = Pt(11)

    def add_heading(text, level=1):
        if level == 1:
            document.add_paragraph(text, style='H1Custom')
        elif level == 2:
            document.add_paragraph(text, style='H2Custom')
        elif level == 3:
            document.add_paragraph(text, style='H3Custom')
        else:
            document.add_heading(text, level=level)

    def add_paragraph(text, style='Normal', alignment=None):
        p = document.add_paragraph(text, style=style)
        if alignment:
            p.alignment = alignment
        return p

    def add_section(section_name, section_data):
        add_heading(section_name, level=1)

        if isinstance(section_data, str):
            add_paragraph(section_data)
        elif isinstance(section_data, list):
            for item in section_data:
                if isinstance(item, str):
                    add_paragraph(item, style='List Bullet')
                elif isinstance(item, dict):
                    add_complex_item(item)
        elif isinstance(section_data, dict):
            add_dictionary_section(section_data)

    def add_complex_item(item_data):
        title = item_data.get('title', item_data.get('position'))
        company = item_data.get('company', item_data.get(
            'institution', item_data.get('organization')))
        dates = item_data.get('dates', item_data.get('years'))
        subtitle_parts = [part for part in [title, company, dates] if part]
        subtitle = " | ".join(subtitle_parts)

        if title:
            add_heading(title, level=2)

        if subtitle:
            add_paragraph(subtitle, style='Subtitle')

        for key in ['responsibilities', 'details', 'description', 'bullets']:
            if key in item_data and isinstance(item_data[key], list):
                for bullet in item_data[key]:
                    add_paragraph(bullet, style='List Bullet')
                break

        for key, value in item_data.items():
            if key not in ['title', 'position', 'company', 'institution', 'organization',
                           'dates', 'years', 'responsibilities', 'details', 'description',
                           'bullets', 'subtitle']:
                if isinstance(value, str):
                    add_paragraph(f"{key.capitalize()}: {value}")
                elif isinstance(value, list):
                    add_paragraph(f"{key.capitalize()}:")
                    for list_item in value:
                        add_paragraph(f"{list_item}", style='List Bullet')

    def add_dictionary_section(data):
        for key, value in data.items():
            if isinstance(value, str):
                add_paragraph(f"{key.capitalize()}: {value}")
            elif isinstance(value, list):
                add_paragraph(f"{key.capitalize()}: {', '.join(value)}")

    # Build Resume
    for section_name, section_data in resume_data.items():
        if section_name == 'name':
            name_paragraph = add_paragraph(
                section_data, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            for run in name_paragraph.runs:
                run.font.size = Pt(30)
        else:
            add_section(section_name.capitalize(), section_data)

    return document


def convert_to_pdf(docx_path, pdf_path):
    """
    Converts a DOCX file to PDF using LibreOffice.
    Requires LibreOffice to be installed on the system.
    """
    try:
        command = [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            os.path.dirname(pdf_path) or ".",
            docx_path,
        ]

        result = subprocess.run(
            command, capture_output=True, text=True, check=True)

        if result.returncode != 0:
            raise Exception(f"Error converting to PDF: {result.stderr}")

        # Rename the output file to match the desired path
        file_name = os.path.splitext(os.path.basename(docx_path))[0]
        created_pdf = os.path.join(os.path.dirname(
            pdf_path) or ".", file_name + ".pdf")
        if created_pdf != pdf_path:
            os.rename(created_pdf, pdf_path)

        return True
    except Exception as e:
        print(f"Error converting to PDF: {str(e)}")
        return False


def generate_resume(resume_data, output_docx, output_pdf=None):
    """
    Main function to generate resume in both DOCX and PDF formats.
    """
    try:
        # Create the DOCX
        document = create_resume_template(resume_data)
        document.save(output_docx)

        # Convert to PDF if requested
        if output_pdf:
            return convert_to_pdf(output_docx, output_pdf)
        return True
    except Exception as e:
        print(f"Error generating resume: {str(e)}")
        return False
